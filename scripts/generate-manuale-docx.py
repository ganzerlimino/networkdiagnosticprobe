#!/usr/bin/env python3
"""Generate docs/NDP-MANUALE-COMPLETO.docx from Markdown sources."""

from __future__ import annotations

import re
import sys
from datetime import date
from pathlib import Path

REPO = Path(__file__).resolve().parents[1]
DOCS = REPO / "docs"
OUTPUT = DOCS / "NDP-MANUALE-COMPLETO.docx"
VERSION = "0.23.1"

SOURCES: list[tuple[str, Path]] = [
    ("Documentazione tecnica e operativa", DOCS / "DOCUMENTAZIONE.md"),
    ("Accesso da telefono e hotspot", DOCS / "MANUALE-ACCESSO-TELEFONO.md"),
    ("Configurazione temi colori", DOCS / "CONFIGURAZIONE-TEMI.md"),
]


def _require_docx():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as exc:
        raise SystemExit("Install python-docx: pip install python-docx") from exc
    return Document, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, qn, Cm, Pt, RGBColor


def _set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(_qn("w:fill"), fill)
    tc_pr.append(shd)


def _add_cover(doc, Document, WD_ALIGN_PARAGRAPH, Pt, RGBColor) -> None:
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Network Diagnostic Probe")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Manuale completo — installazione, uso sul campo, Web UI, OT, config")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ver.add_run(f"Versione software {VERSION}  ·  {date.today():%d/%m/%Y}")
    r.font.size = Pt(11)
    r.italic = True

    doc.add_page_break()


def _parse_table_lines(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    header = [c.strip() for c in lines[0].strip("|").split("|")]
    rows: list[list[str]] = []
    for line in lines[2:]:
        if not line.strip():
            continue
        rows.append([c.strip() for c in line.strip("|").split("|")])
    return header, rows


def _add_table(doc, header: list[str], rows: list[list[str]], Pt, RGBColor, qn) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(header):
        hdr_cells[i].text = text
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        _set_cell_shading(hdr_cells[i], "E8F4FC")
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, text in enumerate(row):
            if c_idx < len(cells):
                cells[c_idx].text = text
                for p in cells[c_idx].paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)
    doc.add_paragraph()


def _add_codeblock(doc, lines: list[str], Pt, Cm) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    text = "\n".join(lines)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(8)


def _convert_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


def _append_markdown(doc, path: Path, section_title: str, Pt, RGBColor, qn, Cm) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not path.is_file():
        doc.add_heading(section_title, level=1)
        doc.add_paragraph(f"(File mancante: {path.name})")
        return

    doc.add_heading(section_title, level=1)
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []
    table_buf: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                _add_codeblock(doc, code_buf, Pt, Cm)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_buf.append(line)
            i += 1
            if i < len(lines) and lines[i].strip().startswith("|"):
                continue
            if len(table_buf) >= 2:
                header, rows = _parse_table_lines(table_buf)
                _add_table(doc, header, rows, Pt, RGBColor, qn)
            table_buf = []
            continue
        table_buf = []

        if stripped in {"---", "***", "___"}:
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(_convert_inline(stripped[2:]), level=2)
        elif stripped.startswith("## "):
            doc.add_heading(_convert_inline(stripped[3:]), level=3)
        elif stripped.startswith("### "):
            doc.add_heading(_convert_inline(stripped[4:]), level=4)
        elif stripped.startswith("#### "):
            p = doc.add_paragraph()
            r = p.add_run(_convert_inline(stripped[5:]))
            r.bold = True
            r.font.size = Pt(11)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(_convert_inline(stripped[2:]), style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            doc.add_paragraph(_convert_inline(re.sub(r"^\d+\.\s", "", stripped)), style="List Number")
        elif stripped.startswith(">"):
            p = doc.add_paragraph(_convert_inline(stripped.lstrip("> ")))
            p.paragraph_format.left_indent = Cm(0.75)
            for run in p.runs:
                run.italic = True
        elif stripped:
            p = doc.add_paragraph(_convert_inline(stripped))
            for run in p.runs:
                run.font.size = Pt(11)

        i += 1

    doc.add_page_break()


def generate(output: Path = OUTPUT) -> Path:
    Document, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, qn, Cm, Pt, RGBColor = _require_docx()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    _add_cover(doc, Document, WD_ALIGN_PARAGRAPH, Pt, RGBColor)

    doc.add_heading("Indice del manuale", level=1)
    for idx, (title, src) in enumerate(SOURCES, start=1):
        doc.add_paragraph(f"{idx}. {title} ({src.name})", style="List Number")
    doc.add_page_break()

    for title, src in SOURCES:
        _append_markdown(doc, src, title, Pt, RGBColor, qn, Cm)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    return output


def main() -> int:
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else OUTPUT
    path = generate(out)
    print(f"Written {path} ({path.stat().st_size // 1024} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
